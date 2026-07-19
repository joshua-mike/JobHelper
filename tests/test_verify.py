"""Offline tests for post-render DOCX verification (ITEM-8).

Round-trips a known content dict through build_resume and verifies the saved
artifact (not the in-memory dict). Structural checks hard-fail; model-behavior
checks (coverage / frequency cap) only warn.

Run:  python tests/test_verify.py
"""
from __future__ import annotations

import copy
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from docx import Document  # noqa: E402
from docx.shared import RGBColor  # noqa: E402

from jobhelper.tailor import verify as V  # noqa: E402
from jobhelper.tailor.resume_docx import build_resume  # noqa: E402


def check(cond, msg):
    print(f"  [{'PASS' if cond else 'FAIL'}] {msg}")
    if not cond:
        raise AssertionError(msg)


CONTENT = {
    "name": "Jane Doe",
    "email": "j@x.com",
    "phone": "555-0100",
    "location": "Remote (US)",
    "links": ["https://li/jane"],
    "summary": "Backend engineer building reliable C# services.",
    "skills": ["C#", ".NET", "PostgreSQL", "AWS"],
    "experience": [
        {"company": "Acme Cloud", "title": "Senior Software Engineer",
         "location": "Remote", "start": "March 2021", "end": "Present",
         "bullets": ["Rebuilt metering pipeline, 6h to 40min."]},
        {"company": "Beta Analytics", "title": "Software Engineer",
         "location": "Austin, TX", "start": "January 2019", "end": "February 2021",
         "bullets": ["Built an ETL framework used by 4 teams."]},
    ],
    "education": [{"institution": "State U", "degree": "B.S.", "field": "CS",
                   "grad": "May 2018"}],
    "certifications": ["CompTIA Security+"],
}


def render(content, td, name="r.docx"):
    path = Path(td) / name
    build_resume(content, path)
    return path


def test_renderer_shape():
    print("== renderer: PROFESSIONAL SUMMARY + 3-line job entry ==")
    with tempfile.TemporaryDirectory() as td:
        path = render(CONTENT, td)
        lines = [p.text for p in Document(str(path)).paragraphs]
    check("PROFESSIONAL SUMMARY" in lines, "heading is PROFESSIONAL SUMMARY")
    i = lines.index("Senior Software Engineer")
    check(lines[i + 1] == "Acme Cloud — Remote", "line 2 is Company — Location")
    check(lines[i + 2] == "March 2021 – Present", "line 3 is the date range")


def test_round_trip_passes():
    print("== round trip: clean render verifies ==")
    with tempfile.TemporaryDirectory() as td:
        path = render(CONTENT, td)
        failures = V.structural_failures(path, CONTENT)
    check(failures == [], f"no structural failures ({failures})")


def test_dropped_daterange_fails():
    print("== mutation: dropped date range ==")
    broken = copy.deepcopy(CONTENT)
    broken["experience"][0]["start"] = ""
    broken["experience"][0]["end"] = ""
    with tempfile.TemporaryDirectory() as td:
        path = render(broken, td)
        failures = V.structural_failures(path, CONTENT)
    check(any("March 2021" in f for f in failures),
          f"missing date range detected ({failures})")


def test_blank_contact_fails():
    print("== mutation: blanked contact ==")
    broken = copy.deepcopy(CONTENT)
    broken["email"] = ""
    broken["phone"] = ""
    with tempfile.TemporaryDirectory() as td:
        path = render(broken, td)
        failures = V.structural_failures(path, CONTENT)
    check(any("contact" in f.lower() for f in failures),
          f"missing contact detected ({failures})")


def test_out_of_order_fails():
    print("== mutation: experience out of order ==")
    swapped = copy.deepcopy(CONTENT)
    swapped["experience"].reverse()
    with tempfile.TemporaryDirectory() as td:
        path = render(swapped, td)
        failures = V.structural_failures(path, CONTENT)
    check(len(failures) > 0, f"order violation detected ({failures})")


def test_missing_heading_fails():
    print("== mutation: dropped section ==")
    broken = copy.deepcopy(CONTENT)
    broken["certifications"] = []
    with tempfile.TemporaryDirectory() as td:
        path = render(broken, td)
        failures = V.structural_failures(path, CONTENT)
    check(any("CERTIFICATIONS" in f for f in failures),
          f"missing heading detected ({failures})")


def test_hidden_runs_fail():
    print("== hidden text scan ==")
    with tempfile.TemporaryDirectory() as td:
        path = render(CONTENT, td)
        doc = Document(str(path))
        r = doc.add_paragraph().add_run("kubernetes kubernetes kubernetes")
        r.font.hidden = True
        doc.save(str(path))
        failures = V.structural_failures(path, CONTENT)
        check(any("hidden" in f.lower() for f in failures),
              f"hidden run detected ({failures})")

        path2 = render(CONTENT, td, "white.docx")
        doc2 = Document(str(path2))
        r2 = doc2.add_paragraph().add_run("terraform terraform")
        r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        doc2.save(str(path2))
        failures2 = V.structural_failures(path2, CONTENT)
        check(any("hidden" in f.lower() for f in failures2),
              f"white-on-white run detected ({failures2})")


GROUPED_CONTENT = {
    **copy.deepcopy(CONTENT),
    "credentials": "Active TS/SCI Clearance  |  CompTIA Security+",
    "skill_groups": [
        {"label": "Languages", "skills": ["C#", ".NET"]},
        {"label": "Data", "skills": ["PostgreSQL"]},
        {"label": None, "skills": ["AWS"]},
    ],
}


def test_grouped_render_and_verify():
    print("== ITEM-13: credentials line + grouped skills render & verify ==")
    with tempfile.TemporaryDirectory() as td:
        path = render(GROUPED_CONTENT, td)
        lines = [p.text for p in Document(str(path)).paragraphs if p.text.strip()]
        failures = V.structural_failures(path, GROUPED_CONTENT)
    check(lines[2] == GROUPED_CONTENT["credentials"],
          "credentials is the 3rd non-empty body line (under contact)")
    check("Languages: C#, .NET" in lines, "grouped line 'Languages: C#, .NET'")
    check("Data: PostgreSQL" in lines, "grouped line 'Data: PostgreSQL'")
    check("AWS" in lines, "ungrouped skills render as a plain line")
    check("C#, .NET, PostgreSQL, AWS" not in lines,
          "flat skills line replaced by grouped lines")
    check(failures == [], f"grouped round trip verifies clean ({failures})")


def test_missing_credentials_or_group_fails():
    print("== ITEM-13 mutation: flat render vs grouped expectation ==")
    with tempfile.TemporaryDirectory() as td:
        path = render(CONTENT, td)  # no credentials, flat skills
        failures = V.structural_failures(path, GROUPED_CONTENT)
    check(any("credentials" in f for f in failures),
          f"missing credentials line detected ({failures})")
    check(any("skills group line missing: Languages" in f for f in failures),
          f"missing group line detected ({failures})")


def test_group_skill_dropped_fails():
    print("== ITEM-13 mutation: skill dropped from its group line ==")
    broken = copy.deepcopy(GROUPED_CONTENT)
    broken["skill_groups"][0]["skills"] = ["C#"]  # .NET vanished
    with tempfile.TemporaryDirectory() as td:
        path = render(broken, td)
        failures = V.structural_failures(path, GROUPED_CONTENT)
    check(any("skill missing from group" in f and ".NET" in f for f in failures),
          f"dropped group member detected ({failures})")


TABLE = [
    {"term": "C#", "category": "hard_skill", "required": True, "variants": []},
    {"term": "Kubernetes", "category": "hard_skill", "required": True,
     "variants": ["K8s"]},
    {"term": ".NET", "category": "hard_skill", "required": False, "variants": []},
]


def test_ats_report():
    print("== build_ats_report ==")
    text = "C# services on .NET. " + "C# " * 5
    report = V.build_ats_report(TABLE, text, missing_required=["Kubernetes"])
    check(report["coverage"]["required_present"] == 1
          and report["coverage"]["required_total"] == 2,
          "coverage 1/2 required")
    check(report["coverage"]["missing"] == ["Kubernetes"], "missing recorded")
    check(report["missing_required"] == ["Kubernetes"],
          "tailor-flagged missing_required kept")
    check(any("C#" in w for w in report["warnings"]),
          f"frequency cap warning for C# (>4 hits) ({report['warnings']})")
    check(report["keyword_table"] is TABLE or report["keyword_table"] == TABLE,
          "keyword table stored in blob")

    bare = V.build_ats_report(None, text, missing_required=None)
    check(bare.get("coverage") is None and bare["warnings"] == [],
          "no-table report has no coverage and no warnings")


def main() -> int:
    test_renderer_shape()
    test_round_trip_passes()
    test_dropped_daterange_fails()
    test_blank_contact_fails()
    test_out_of_order_fails()
    test_missing_heading_fails()
    test_hidden_runs_fail()
    test_grouped_render_and_verify()
    test_missing_credentials_or_group_fails()
    test_group_skill_dropped_fails()
    test_ats_report()
    print("\nALL VERIFY CHECKS PASSED")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
