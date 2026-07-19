"""Offline wiring test for the LLM path — no API key, no network.

Uses a fake LLM that returns canned structured output to prove:
  1. tailor assembly drops INVENTED skills and keeps only profile skills,
  2. company/title/dates are preserved verbatim (LLM only supplies bullets),
  3. the LLM judge fields flow into the digest,
  4. the ATS-safe .docx renders with zero tables.

Run:  python tests/test_tailor_wiring.py
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from docx import Document  # noqa: E402

from jobhelper.tailor import tailor as T  # noqa: E402
from jobhelper.tailor.resume_docx import build_resume  # noqa: E402
from jobhelper.digest.digest import render_digest  # noqa: E402
from jobhelper.util import RESUME_DIR  # noqa: E402

PROFILE = {
    "identity": {"full_name": "Jane Doe", "email": "j@x.com", "phone": "555",
                 "city_state": "Remote (US)", "linkedin_url": "https://li/jane",
                 "work_authorization_status": "US authorized",
                 "credentials_line": "Active Clearance  |  CompTIA Security+",
                 "requires_sponsorship": False, "willing_to_relocate": False,
                 "earliest_start_date": "2 weeks", "notice_period": "2 weeks"},
    "compensation": {"desired_salary_min": 120000, "desired_salary_max": 160000,
                     "currency": "USD", "salary_negotiable": True},
    "summary": "Backend engineer.",
    "work_history": [
        {"company": "Acme Cloud", "title": "Senior Software Engineer",
         "location": "Remote", "start_date": "2021-03", "end_date": "Present",
         "achievements": [{"text": "Rebuilt metering pipeline 6h->40min."},
                          {"text": "Shipped billing API for 12k customers."}]},
        {"company": "Beta Analytics", "title": "Software Engineer",
         "location": "Austin, TX", "start_date": "2019-01", "end_date": "2021-02",
         "achievements": [{"text": "Built an ETL framework used by 4 teams."}]},
    ],
    "education": [{"institution": "State U", "degree": "B.S.", "field": "CS",
                   "grad_date": "2018-05"}],
    "skills": {"hard_skills": [{"name": "Python", "group": "Languages"},
                               {"name": "FastAPI", "group": "Languages"},
                               {"name": "PostgreSQL", "group": "Data"},
                               {"name": "AWS"}],
               "soft_skills": [], "certifications": []},
    "eeo": {"gender": "decline to self-identify"},
}

JOB = {"title": "Senior Backend Engineer", "company": "Globex",
       "location": "Remote", "description_clean": "Python, FastAPI, distributed systems."}


class FakeLLM:
    """Mimics jobhelper.llm.LLM with canned, adversarial output."""
    available = True

    def structured(self, system, user, *, schema, tool_name, model, max_tokens=1024):
        return {
            "summary": "Backend engineer focused on reliable Python services.",
            # 'Rust' and 'Kubernetes' are NOT in the profile — must be dropped.
            "skills_order": [{"skill": "FastAPI"}, {"skill": "Python"},
                             {"skill": "Rust"}, {"skill": "Kubernetes"}],
            "jobs": [
                {"index": 0, "bullets": ["Cut nightly metering batch from 6h to 40min."]},
                # index 1 omitted -> should fall back to profile achievements.
            ],
            "change_notes": ["Reordered skills; emphasized metering work."],
            "missing_required": ["Go"],
        }

    def text(self, system, user, *, model, max_tokens=1024):
        return "Dear Globex team,\nI'd love to help...\nBest, Jane"


def check(cond: bool, msg: str) -> None:
    print(f"  [{'PASS' if cond else 'FAIL'}] {msg}")
    if not cond:
        raise AssertionError(msg)


def main() -> int:
    print("== tailor_resume assembly ==")
    content, notes, missing = T.tailor_resume(FakeLLM(), "fake-model", PROFILE, JOB)
    check(missing == ["Go"], "missing_required flows out of tailor_resume")

    profile_skills = {"python", "fastapi", "postgresql", "aws"}
    out_skills = {s.lower() for s in content["skills"]}
    check("rust" not in out_skills, "invented skill 'Rust' dropped")
    check("kubernetes" not in out_skills, "invented skill 'Kubernetes' dropped")
    check(profile_skills <= out_skills, "all real profile skills retained")
    check(content["skills"][0] == "FastAPI", "skill reordering applied (FastAPI first)")

    exp = content["experience"]
    check(exp[0]["company"] == "Acme Cloud" and exp[0]["title"] == "Senior Software Engineer",
          "job 0 company/title preserved verbatim")
    check(exp[0]["start"] == "March 2021" and exp[0]["end"] == "Present",
          "job 0 dates formatted from profile (not LLM)")
    check(exp[0]["bullets"] == ["Cut nightly metering batch from 6h to 40min."],
          "job 0 uses LLM-reworded bullet")
    check(exp[1]["bullets"] == ["Built an ETL framework used by 4 teams."],
          "job 1 (omitted by LLM) falls back to profile achievements")
    check(content["summary"].startswith("Backend engineer focused"), "summary tailored")
    check(len(notes) >= 1, "change notes captured")

    print("== ITEM-13: credentials + skill groups through tailoring ==")
    check(content["credentials"] == "Active Clearance  |  CompTIA Security+",
          "credentials line passed through untouched")
    sg = content["skill_groups"]
    check([g["label"] for g in sg] == ["Languages", "Data", None],
          f"group order follows the profile ({[g['label'] for g in sg]})")
    check(sg[0]["skills"] == ["FastAPI", "Python"],
          f"within-group order follows tailored relevance ({sg[0]['skills']})")
    check(sg[1]["skills"] == ["PostgreSQL"] and sg[2]["skills"] == ["AWS"],
          "dropped-by-LLM skills still land in their groups")
    all_grouped = [s for g in sg for s in g["skills"]]
    check("Rust" not in all_grouped and "Kubernetes" not in all_grouped,
          "invented skills never reach the groups")

    print("== ATS-safe .docx render ==")
    out = RESUME_DIR / "_selftest.docx"
    build_resume(content, out)
    doc = Document(str(out))
    check(out.exists(), "docx written")
    check(len(doc.tables) == 0, "zero tables (ATS-safe single column)")
    headings = [p.text for p in doc.paragraphs if p.runs and p.runs[0].bold]
    check(any(h == "WORK EXPERIENCE" for h in headings), "whitelisted heading present")
    doc_lines = [p.text for p in doc.paragraphs if p.text.strip()]
    check("Active Clearance  |  CompTIA Security+" in doc_lines,
          "credentials line rendered in the docx")
    check("Languages: FastAPI, Python" in doc_lines,
          "grouped skills line rendered in tailored order")
    out.unlink(missing_ok=True)

    print("== screening answers ==")
    ans = T.screening_answers(PROFILE)
    check(ans["years_of_experience"] > 0, f"years computed ({ans['years_of_experience']})")
    check(ans["requires_sponsorship"] is False, "sponsorship flag from profile")

    print("== digest renders LLM fields ==")
    job_row = {"title": JOB["title"], "company": JOB["company"], "location": "Remote",
               "source": "greenhouse", "url": "https://x/y", "llm_score": 82,
               "llm_rationale": "Strong Python/FastAPI overlap.",
               "llm_musthaves_met": '["Python","FastAPI"]', "llm_missing": '["Go"]',
               "tailored_resume_path": "data/resumes/x.docx",
               "cover_letter_text": "Dear Globex team,\nHello.",
               "change_log": '["Reordered skills"]', "screening_answers": "{}",
               "embed_score": 0.4,
               "ats_report": '{"coverage": {"required_present": 1, '
                             '"required_total": 2, "missing": ["Go"]}, '
                             '"warnings": ["\'Python\' appears 5x document-wide '
                             '(cap 4)"]}'}
    md, path = render_digest([job_row], "testrun", "lexical", llm_on=True)
    check("score 82" in md, "llm_score shown in digest")
    check("Strong Python/FastAPI overlap" in md, "rationale shown")
    check("Meets:" in md and "Gaps:" in md, "met/missing shown")
    check("Cover letter draft" in md, "cover letter shown")
    check("ATS coverage:** 1/2 required · missing: Go" in md,
          "ats coverage line shown")
    check("appears 5x" in md, "ats warning shown")
    path.unlink(missing_ok=True)

    print("\nALL WIRING CHECKS PASSED")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
