"""Bootstrap or refresh config/profile.yaml from an uploaded resume.

Claude-only (LLM.structured with the criteria tailor_model): there is no
regex fallback because a badly parsed profile silently poisons every later
stage (scoring, tailoring, screening). Accepts .docx/.txt/.md — no PDF.

The merge is SECTIONAL: summary / work_history / education / skills and the
identity contact fields are proposed from the resume, while compensation,
eeo, qa_bank and the identity work-authorization/notice fields are never
touched by an import — they're preserved from the existing profile, or
seeded from profile.example.yaml on a fresh clone. Nothing is written here;
the caller previews the proposal and saves via the normal profile PUT.
"""
from __future__ import annotations

import copy
import io
from typing import Any

from ..llm import LLM
from ..util import get_logger

log = get_logger()

ALLOWED_EXTENSIONS = (".docx", ".txt", ".md")
MAX_UPLOAD_BYTES = 5 * 1024 * 1024
MAX_TEXT_CHARS = 60_000

# Identity fields a resume can legitimately provide vs. the ones it can't.
CONTACT_FIELDS = ("full_name", "email", "phone", "city_state",
                  "linkedin_url", "portfolio_url")
PRESERVED_IDENTITY_FIELDS = ("work_authorization_status", "requires_sponsorship",
                             "willing_to_relocate", "earliest_start_date",
                             "notice_period")
RESUME_SECTIONS = ("summary", "work_history", "education", "skills")
PRESERVED_SECTIONS = ("compensation", "eeo", "qa_bank")


# ---- File -> text ----------------------------------------------------------------
def extract_text(filename: str, data: bytes) -> str:
    """Plain text from an upload. Raises ValueError on unsupported/undecodable."""
    name = (filename or "").lower()
    if name.endswith(".docx"):
        text = _docx_text(data)
    elif name.endswith((".txt", ".md")):
        text = data.decode("utf-8", errors="replace")
    else:
        raise ValueError(
            "Unsupported file type — upload a .docx, .txt, or .md resume "
            "(PDF is not supported).")
    text = text.strip()
    if not text:
        raise ValueError("The file contains no readable text.")
    return text[:MAX_TEXT_CHARS]


def _docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - python-docx is a core dep
        raise ValueError("python-docx is not installed.") from exc
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise ValueError(f"Could not read the .docx file: {exc}") from exc
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(c for c in cells if c))
    return "\n".join(parts)


# ---- LLM extraction --------------------------------------------------------------
EXTRACT_INSTRUCTIONS = (
    "You extract structured data from a resume. Use ONLY facts stated in the "
    "resume text — never invent or infer employers, titles, dates, metrics, "
    "skills, or contact details. Dates: YYYY-MM when the resume gives a month, "
    "else YYYY; 'Present' for current roles. Omit optional fields the resume "
    "does not state (do not guess per-skill years). achievements: one entry per "
    "resume bullet with text near-verbatim (light cleanup only) and skills_used "
    "listing only technologies actually named in that bullet."
)

_STR = {"type": "string"}
_STR_LIST = {"type": "array", "items": _STR}

EXTRACT_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "identity": {
            "type": "object",
            "properties": {f: _STR for f in CONTACT_FIELDS},
        },
        "summary": _STR,
        "work_history": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "company": _STR, "title": _STR, "location": _STR,
                    "start_date": _STR, "end_date": _STR,
                    "employment_type": _STR, "summary": _STR,
                    "achievements": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {"text": _STR,
                                           "skills_used": _STR_LIST},
                            "required": ["text"],
                        },
                    },
                },
                "required": ["company", "title"],
            },
        },
        "education": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {"institution": _STR, "degree": _STR,
                               "field": _STR, "grad_date": _STR},
                "required": ["institution"],
            },
        },
        "skills": {
            "type": "object",
            "properties": {
                "hard_skills": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {"name": _STR,
                                       "years": {"type": "number"},
                                       "proficiency": _STR},
                        "required": ["name"],
                    },
                },
                "soft_skills": _STR_LIST,
                "certifications": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {"name": _STR, "issuer": _STR,
                                       "date": _STR, "expiry": _STR},
                        "required": ["name"],
                    },
                },
                "languages": _STR_LIST,
            },
        },
    },
    "required": ["identity", "summary", "work_history", "education", "skills"],
}


def extract_profile(llm: LLM, model: str, resume_text: str) -> dict[str, Any] | None:
    return llm.structured(
        EXTRACT_INSTRUCTIONS,
        f"Resume text:\n\n{resume_text}",
        schema=EXTRACT_SCHEMA,
        tool_name="resume_profile",
        model=model,
        max_tokens=8192,
    )


# ---- Sectional merge --------------------------------------------------------------
def _strip_nones(value: Any) -> Any:
    """Drop None-valued keys so the proposal doesn't write `key: null` noise."""
    if isinstance(value, dict):
        return {k: _strip_nones(v) for k, v in value.items() if v is not None}
    if isinstance(value, list):
        return [_strip_nones(v) for v in value]
    return value


def _norm_work_history(entries: list[dict]) -> list[dict]:
    out = []
    for job in entries:
        job = dict(_strip_nones(job))
        achs = []
        for a in job.get("achievements") or []:
            a = dict(a) if isinstance(a, dict) else {"text": str(a)}
            text = (a.get("text") or "").strip()
            if not text:
                continue
            # Imported metrics start unverified — the human flips the flag once
            # they can defend the number (profile contract).
            achs.append({"text": text,
                         "skills_used": a.get("skills_used") or [],
                         "verified": False})
        job["achievements"] = achs
        out.append(job)
    return out


def sectional_merge(extracted: dict[str, Any], base: dict[str, Any],
                    *, bootstrapped: bool) -> tuple[dict, list[dict[str, str]]]:
    """Overlay resume-derived sections on `base`; never touch preserved ones.

    Returns (proposed_profile, section_notes). `bootstrapped` means base came
    from profile.example.yaml (no profile.yaml yet), which the notes call out
    as 'seeded' so the UI can tell the user to review those placeholders.
    """
    proposed = copy.deepcopy(base)
    notes: list[dict[str, str]] = []
    kept = "seeded" if bootstrapped else "preserved"
    kept_detail = ("placeholder values seeded from profile.example.yaml — review before relying on them"
                   if bootstrapped else "kept from your existing profile")

    ident = proposed.setdefault("identity", {})
    new_ident = extracted.get("identity") or {}
    imported_fields = []
    for f in CONTACT_FIELDS:
        val = (new_ident.get(f) or "").strip()
        if val:
            ident[f] = val
            imported_fields.append(f)
    notes.append({
        "section": "identity", "action": "imported",
        "detail": (f"contact fields from resume: {', '.join(imported_fields)}"
                   if imported_fields else "no contact details found in resume")
                  + "; work-authorization and notice fields " + kept})

    for key in RESUME_SECTIONS:
        val = extracted.get(key)
        if key == "work_history" and val:
            val = _norm_work_history(val)
        elif val is not None:
            val = _strip_nones(val)
        if val:
            proposed[key] = val
            detail = {
                "summary": "professional summary from resume",
                "work_history": f"{len(val)} position(s) from resume — achievements marked unverified until you confirm the numbers",
                "education": f"{len(val)} entr{'y' if len(val) == 1 else 'ies'} from resume",
                "skills": "skills from resume",
            }[key]
            notes.append({"section": key, "action": "imported", "detail": detail})
        else:
            notes.append({"section": key, "action": kept,
                          "detail": f"nothing found in resume; {kept_detail}"})

    for key in PRESERVED_SECTIONS:
        notes.append({"section": key, "action": kept, "detail": kept_detail})
    return proposed, notes
