"""Render an ATS-safe single-column .docx resume from a ResumeContent dict.

ATS rules baked in: single column (no tables/text boxes/columns), contact info in
the BODY (never header/footer), system font, whitelisted section headings,
standard bullets. See research notes in README.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor

FONT = "Calibri"


def _heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = FONT
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _line(doc: Document, text: str, *, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = FONT


def build_resume(content: dict[str, Any], path: Path) -> Path:
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(54)   # 0.75"
        section.left_margin = section.right_margin = Pt(72)   # 1"

    # --- Contact block in the BODY, top-left ---
    name_p = doc.add_paragraph()
    name_run = name_p.add_run(content.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.name = FONT

    contact_bits = [content.get("email"), content.get("phone"),
                    content.get("location"), *content.get("links", [])]
    contact = "  |  ".join(b for b in contact_bits if b)
    if contact:
        _line(doc, contact, size=10)

    # Credentials line (clearance + certs) directly under contact — recruiter
    # boolean searches hit these literal tokens, so they belong at the top.
    if content.get("credentials"):
        _line(doc, content["credentials"], size=10)

    # --- Summary ---
    if content.get("summary"):
        _heading(doc, "Professional Summary")
        _line(doc, content["summary"])

    # --- Skills ---
    # Grouped: one plain "Label: a, b, c" line per group (parse-safe, no tables).
    # Ungrouped fallback: the flat comma line (bare tokens: NER-friendly).
    groups = content.get("skill_groups") or []
    skills = content.get("skills", [])
    if groups or skills:
        _heading(doc, "Skills")
    if groups:
        for g in groups:
            names = [s for s in (g.get("skills") or []) if s]
            if not names:
                continue
            label = g.get("label")
            if label:
                p = doc.add_paragraph()
                r1 = p.add_run(f"{label}: ")
                r1.bold = True
                r1.font.size = Pt(11)
                r1.font.name = FONT
                r2 = p.add_run(", ".join(names))
                r2.font.size = Pt(11)
                r2.font.name = FONT
            else:
                _line(doc, ", ".join(names))
    elif skills:
        _line(doc, ", ".join(skills))

    # --- Work Experience ---
    exp = content.get("experience", [])
    if exp:
        _heading(doc, "Work Experience")
        # 3 lines per entry — parsers key off line structure, and a combined
        # "Title — Company" line is the exact pattern they mis-split.
        for job in exp:
            _line(doc, job.get("title", ""), bold=True)
            company = " — ".join(b for b in [job.get("company"),
                                             job.get("location")] if b)
            if company:
                _line(doc, company)
            dates = daterange(job)
            if dates:
                _line(doc, dates, size=10)
            for bullet in job.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                r = bp.add_run(bullet)
                r.font.size = Pt(11)
                r.font.name = FONT

    # --- Education ---
    edu = content.get("education", [])
    if edu:
        _heading(doc, "Education")
        for e in edu:
            deg = " ".join(x for x in [e.get("degree"), e.get("field")] if x)
            line = f"{deg} — {e.get('institution', '')}".strip(" —")
            _line(doc, line, bold=True)
            if e.get("grad"):
                _line(doc, e["grad"], size=10)

    # --- Certifications ---
    certs = content.get("certifications", [])
    if certs:
        _heading(doc, "Certifications")
        _line(doc, ", ".join(certs))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def daterange(job: dict[str, Any]) -> str:
    """Public: verify.py checks the artifact against this exact rendering."""
    start, end = job.get("start", ""), job.get("end", "")
    if start or end:
        return f"{start} – {end}".strip(" –")
    return ""
