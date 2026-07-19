"""Post-render verification of the saved resume DOCX.

Verifies the artifact, not the in-memory dict: text is re-extracted from the
file with python-docx the way an ATS parser would read it. Structural failures
hard-fail the job (pipeline marks status='error'); model-behavior findings
(keyword coverage, frequency cap) only warn. No auto-retry in v1.
"""
from __future__ import annotations

import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.shared import RGBColor

from .keywords import FREQUENCY_CAP, coverage
from .resume_docx import daterange

_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Metric-once rule (ITEM-14): a metric may appear in the summary plus one
# bullet — the same 2-placement budget keywords get — never more.
METRIC_CAP = 2

# Metrics are percentages (99.9%, 99%+), N+ counts (10,000+), and dollar
# amounts ($1.2M). Deliberately NOT bare numbers: 'OAuth 2.0', 'RFC 6238',
# and '8570/8140' are spec tokens, not achievements.
_METRIC_RE = re.compile(
    r"\$\d[\d,]*(?:\.\d+)?[KMBkmb]?|\d[\d,]*(?:\.\d+)?%\+?|\d[\d,]*\+")

# A distinctive achievement "survives" when most of its content words are
# still present — light rewording passes, dropping or gutting it does not.
_SURVIVAL_THRESHOLD = 0.6


def _survives(achievement: str, text: str) -> bool:
    words = set(re.findall(r"[a-z0-9][a-z0-9.+#/-]{3,}", achievement.lower()))
    if not words:
        return False
    text_l = text.lower()
    return sum(1 for w in words if w in text_l) / len(words) >= _SURVIVAL_THRESHOLD


def extract_docx_text(path: Path | str) -> str:
    """Plain text as a parser would see it: one line per paragraph, in order."""
    return "\n".join(p.text for p in Document(str(path)).paragraphs)


def _expected_headings(content: dict) -> list[str]:
    wanted = [("summary", "PROFESSIONAL SUMMARY"), ("skills", "SKILLS"),
              ("experience", "WORK EXPERIENCE"), ("education", "EDUCATION"),
              ("certifications", "CERTIFICATIONS")]
    return [h for key, h in wanted if content.get(key)]


def structural_failures(path: Path | str, content: dict) -> list[str]:
    """Hard-fail checks: every failure means the artifact is unusable as-is."""
    doc = Document(str(path))
    lines = [p.text for p in doc.paragraphs]
    text = "\n".join(lines)
    failures: list[str] = []

    # Contact info must sit in the first 5 (non-empty) lines of the BODY —
    # parsers skip header/footer regions, and we never write those anyway.
    head = "\n".join([ln for ln in lines if ln.strip()][:5])
    for label in ("name", "email", "phone"):
        val = content.get(label)
        if val and val not in head:
            failures.append(f"contact {label} not in first 5 lines: {val!r}")

    # Credentials line (clearance/certs) belongs in the same head window.
    cred = content.get("credentials")
    if cred and cred not in head:
        failures.append(f"credentials line not in first 5 lines: {cred!r}")

    # Grouped skills: every group renders as a "Label: a, b, c" line carrying
    # ALL of its skills — a dropped line silently loses those tokens.
    for g in content.get("skill_groups") or []:
        label = g.get("label")
        names = [s for s in (g.get("skills") or []) if s]
        if not label or not names:
            continue
        line = next((ln for ln in lines if ln.startswith(f"{label}:")), None)
        if line is None:
            failures.append(f"skills group line missing: {label}")
            continue
        for s in names:
            if s not in line:
                failures.append(f"skill missing from group {label!r}: {s!r}")

    for h in _expected_headings(content):
        if h not in lines:
            failures.append(f"section heading missing: {h}")

    # Every title, company, and date range from the content dict, IN ORDER.
    expected: list[str] = []
    for job in content.get("experience", []):
        for part in (job.get("title"), job.get("company"), daterange(job)):
            if part:
                expected.append(part)
    pos = 0
    for part in expected:
        found = text.find(part, pos)
        if found < 0:
            kind = "out of order" if part in text else "missing"
            failures.append(f"experience field {kind}: {part!r}")
        else:
            pos = found + len(part)

    # We never write hidden or white runs; the scan makes that guarantee
    # explicit on the artifact (ATSs flag hidden text as manipulation).
    for p in doc.paragraphs:
        for r in p.runs:
            rgb = None
            try:
                if r.font.color is not None and r.font.color.type is not None:
                    rgb = r.font.color.rgb
            except (AttributeError, ValueError):
                rgb = None
            if r.font.hidden or rgb == _WHITE:
                failures.append(f"hidden/white text run: {r.text[:40]!r}")

    return failures


def build_ats_report(keyword_table: list[dict] | None, text: str,
                     missing_required: list[str] | None,
                     distinctive_texts: list[str] | None = None) -> dict:
    """Assemble the jobs.ats_report blob (coverage + keyword warnings need a
    table; metric and distinctive warnings only need the text)."""
    report: dict = {
        "keyword_table": keyword_table,
        "coverage": None,
        "missing_required": list(missing_required or []),
        "warnings": [],
    }
    if keyword_table:
        cov = coverage(text, keyword_table)
        report["coverage"] = {"required_present": cov["required_present"],
                              "required_total": cov["required_total"],
                              "missing": cov["missing"]}
        for term, n in cov["hits"].items():
            if n > FREQUENCY_CAP:
                report["warnings"].append(
                    f"'{term}' appears {n}x document-wide (cap {FREQUENCY_CAP})")

    # Metric-once rule (ITEM-14): repeated metrics read as padding to humans
    # and AI summarizers alike.
    for token, n in Counter(_METRIC_RE.findall(text)).items():
        if n > METRIC_CAP:
            report["warnings"].append(
                f"metric '{token}' appears {n}x document-wide (cap {METRIC_CAP})")

    # Distinctive-survival check (FR-5.2): if every flagged achievement was
    # reworded away, the resume summarizes like everyone else's.
    if distinctive_texts and not any(_survives(t, text) for t in distinctive_texts):
        report["warnings"].append(
            "no distinctive achievement survived tailoring (FR-5.2) — "
            "flagged specifics were dropped or reworded beyond recognition")
    return report
